#!/usr/bin/env python3
"""
Document Parser for Regulatory Standards
Parses PDF and DOCX files to extract standards, sections, and clauses
"""

import re
import logging
from typing import Dict, List, Optional, Tuple
from pathlib import Path

try:
    import pdfplumber
except ImportError:
    pdfplumber = None

try:
    from docx import Document
except ImportError:
    Document = None

from bs4 import BeautifulSoup

logger = logging.getLogger(__name__)


class StandardParser:
    """Parser for regulatory standards documents"""

    def __init__(self):
        self.clause_pattern = re.compile(r'^\d+(\.\d+)*\.?\s')  # Matches "3.2.1" or "3.2.1."
        self.section_pattern = re.compile(r'^(Chapter|Section|Part)\s+(\d+)', re.IGNORECASE)

    def parse_document(self, file_path: str, mime_type: str) -> Dict:
        """
        Parse a document and extract structure

        Args:
            file_path: Path to the document
            mime_type: MIME type of the document

        Returns:
            Dictionary with parsed structure
        """
        try:
            if 'pdf' in mime_type:
                return self.parse_pdf(file_path)
            elif 'document' in mime_type or 'docx' in mime_type:
                return self.parse_docx(file_path)
            elif 'text' in mime_type or 'plain' in mime_type:
                return self.parse_text(file_path)
            else:
                logger.warning(f"Unsupported mime type: {mime_type}")
                return self._empty_structure()

        except Exception as e:
            logger.error(f"Error parsing document: {e}")
            return self._empty_structure()

    def parse_pdf(self, file_path: str) -> Dict:
        """Parse PDF document"""
        if pdfplumber is None:
            logger.error("pdfplumber not installed")
            return self._empty_structure()

        try:
            sections = []
            clauses = []
            full_text = []

            with pdfplumber.open(file_path) as pdf:
                for page_num, page in enumerate(pdf.pages, 1):
                    text = page.extract_text()
                    if text:
                        full_text.append(text)
                        # Extract sections and clauses from page
                        page_sections, page_clauses = self._extract_structure_from_text(
                            text, page_num
                        )
                        sections.extend(page_sections)
                        clauses.extend(page_clauses)

            # Post-process to organize hierarchy
            organized = self._organize_hierarchy(sections, clauses)

            return {
                'type': 'pdf',
                'total_pages': len(pdf.pages) if hasattr(pdf, 'pages') else 0,
                'full_text': '\n'.join(full_text),
                'sections': organized['sections'],
                'clauses': organized['clauses'],
                'metadata': self._extract_metadata(full_text)
            }

        except Exception as e:
            logger.error(f"Error parsing PDF: {e}")
            return self._empty_structure()

    def parse_docx(self, file_path: str) -> Dict:
        """Parse DOCX document"""
        if Document is None:
            logger.error("python-docx not installed")
            return self._empty_structure()

        try:
            doc = Document(file_path)
            sections = []
            clauses = []
            full_text = []

            current_section = None

            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    continue

                full_text.append(text)

                # Check if it's a section heading
                if para.style.name.startswith('Heading'):
                    level = int(para.style.name.replace('Heading ', '')) if para.style.name != 'Heading' else 1
                    section = {
                        'number': self._extract_number(text),
                        'title': self._clean_title(text),
                        'level': level,
                        'content': text
                    }
                    sections.append(section)
                    current_section = section['number']

                # Check if it's a clause
                elif self.clause_pattern.match(text):
                    clause = self._parse_clause(text, current_section)
                    if clause:
                        clauses.append(clause)

            organized = self._organize_hierarchy(sections, clauses)

            return {
                'type': 'docx',
                'total_paragraphs': len(doc.paragraphs),
                'full_text': '\n'.join(full_text),
                'sections': organized['sections'],
                'clauses': organized['clauses'],
                'metadata': self._extract_metadata(full_text)
            }

        except Exception as e:
            logger.error(f"Error parsing DOCX: {e}")
            return self._empty_structure()

    def parse_text(self, file_path: str) -> Dict:
        """Parse plain text document"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                full_text = f.read()

            lines = full_text.split('\n')
            sections, clauses = self._extract_structure_from_text(full_text)
            organized = self._organize_hierarchy(sections, clauses)

            return {
                'type': 'text',
                'total_lines': len(lines),
                'full_text': full_text,
                'sections': organized['sections'],
                'clauses': organized['clauses'],
                'metadata': self._extract_metadata([full_text])
            }

        except Exception as e:
            logger.error(f"Error parsing text: {e}")
            return self._empty_structure()

    def _extract_structure_from_text(self, text: str, page_num: Optional[int] = None) -> Tuple[List, List]:
        """Extract sections and clauses from text"""
        sections = []
        clauses = []

        lines = text.split('\n')
        current_section = None

        for line_num, line in enumerate(lines):
            line = line.strip()
            if not line:
                continue

            # Check for section headings
            section_match = self.section_pattern.match(line)
            if section_match:
                section = {
                    'number': section_match.group(2),
                    'title': line,
                    'level': 1,
                    'content': line,
                    'page': page_num
                }
                sections.append(section)
                current_section = section['number']
                continue

            # Check for clauses (numbered paragraphs)
            if self.clause_pattern.match(line):
                clause = self._parse_clause(line, current_section, page_num)
                if clause:
                    clauses.append(clause)

        return sections, clauses

    def _parse_clause(self, text: str, section_number: Optional[str] = None,
                     page_num: Optional[int] = None) -> Optional[Dict]:
        """Parse a clause from text"""
        try:
            # Extract clause number
            match = self.clause_pattern.match(text)
            if not match:
                return None

            # Get the matched number part
            number_part = text[:match.end()].strip().rstrip('.')
            content = text[match.end():].strip()

            # Split into title and content if there's a colon or period
            title = None
            if ':' in content[:100]:  # Check first 100 chars for title
                parts = content.split(':', 1)
                title = parts[0].strip()
                content = parts[1].strip() if len(parts) > 1 else content
            elif '.' in content[:50]:
                parts = content.split('.', 1)
                if len(parts[0]) < 50:  # Short enough to be a title
                    title = parts[0].strip()
                    content = parts[1].strip() if len(parts) > 1 else content

            return {
                'number': number_part,
                'title': title,
                'content': content,
                'section_number': section_number,
                'page': page_num,
                'category': self._infer_category(content),
                'importance': self._infer_importance(content)
            }

        except Exception as e:
            logger.error(f"Error parsing clause: {e}")
            return None

    def _organize_hierarchy(self, sections: List[Dict], clauses: List[Dict]) -> Dict:
        """Organize sections and clauses into hierarchy"""
        # Group clauses by section
        clause_by_section = {}
        for clause in clauses:
            section = clause.get('section_number', 'general')
            if section not in clause_by_section:
                clause_by_section[section] = []
            clause_by_section[section].append(clause)

        # Add clause counts to sections
        for section in sections:
            section_num = section.get('number')
            section['clause_count'] = len(clause_by_section.get(section_num, []))

        return {
            'sections': sections,
            'clauses': clauses,
            'clause_by_section': clause_by_section
        }

    def _extract_number(self, text: str) -> str:
        """Extract number from text"""
        match = re.match(r'^(\d+(\.\d+)*)', text)
        return match.group(1) if match else ''

    def _clean_title(self, text: str) -> str:
        """Clean title text"""
        # Remove leading numbers and punctuation
        cleaned = re.sub(r'^[\d\.\s]+', '', text)
        return cleaned.strip()

    def _infer_category(self, content: str) -> str:
        """Infer category from content"""
        content_lower = content.lower()

        if any(word in content_lower for word in ['safety', 'hazard', 'protection', 'ppe']):
            return 'Safety'
        elif any(word in content_lower for word in ['equipment', 'dehumidifier', 'air mover']):
            return 'Equipment'
        elif any(word in content_lower for word in ['document', 'record', 'report', 'log']):
            return 'Documentation'
        elif any(word in content_lower for word in ['procedure', 'process', 'method', 'step']):
            return 'Process'
        elif any(word in content_lower for word in ['moisture', 'humidity', 'drying', 'water']):
            return 'Moisture Control'
        elif any(word in content_lower for word in ['inspect', 'assess', 'evaluat', 'test']):
            return 'Inspection'
        else:
            return 'General'

    def _infer_importance(self, content: str) -> str:
        """Infer importance level from content"""
        content_lower = content.lower()

        if any(word in content_lower for word in ['must', 'shall', 'required', 'mandatory']):
            return 'REQUIRED'
        elif any(word in content_lower for word in ['critical', 'danger', 'warning', 'hazard']):
            return 'CRITICAL'
        elif any(word in content_lower for word in ['should', 'recommend', 'suggest']):
            return 'RECOMMENDED'
        elif any(word in content_lower for word in ['may', 'optional', 'consider']):
            return 'OPTIONAL'
        else:
            return 'STANDARD'

    def _extract_metadata(self, text_blocks: List[str]) -> Dict:
        """Extract metadata from document"""
        # Combine first few blocks for metadata extraction
        combined = ' '.join(text_blocks[:5]) if text_blocks else ''

        metadata = {
            'edition': None,
            'version': None,
            'publication_year': None,
            'publisher': None
        }

        # Try to extract edition
        edition_match = re.search(r'(\d+(?:st|nd|rd|th)?\s+Edition)', combined, re.IGNORECASE)
        if edition_match:
            metadata['edition'] = edition_match.group(1)

        # Try to extract year
        year_match = re.search(r'(20\d{2})', combined)
        if year_match:
            metadata['publication_year'] = year_match.group(1)

        # Try to extract publisher
        if 'IICRC' in combined:
            metadata['publisher'] = 'IICRC'
        elif 'Standards Australia' in combined:
            metadata['publisher'] = 'Standards Australia'

        return metadata

    def _empty_structure(self) -> Dict:
        """Return empty structure"""
        return {
            'type': 'unknown',
            'full_text': '',
            'sections': [],
            'clauses': [],
            'metadata': {}
        }


# Utility function for external use
def parse_standard_document(file_path: str, mime_type: str = 'application/pdf') -> Dict:
    """
    Parse a standard document and return structured data

    Args:
        file_path: Path to the document
        mime_type: MIME type (default: application/pdf)

    Returns:
        Parsed structure dictionary
    """
    parser = StandardParser()
    return parser.parse_document(file_path, mime_type)
